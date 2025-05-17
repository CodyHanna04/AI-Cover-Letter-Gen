import os
import io
from dotenv import load_dotenv
import openai
import streamlit as st

# for PDF/DOCX parsing
from PyPDF2 import PdfReader
import docx

# — Page config for a more polished appearance —
st.set_page_config(
    page_title="AI Cover Letter Maker",
    page_icon="✉️",
    layout="centered",
    initial_sidebar_state="auto"
)

load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

# — Sidebar branding/info —
with st.sidebar:
    st.markdown("## 📄 AI Cover Letter Maker")
    st.write(
        "Generate tailored cover letters in seconds.\n\n"
        "Upload your resume and fill out the job details below."
    )
    st.markdown("---")
    st.write("Built with Streamlit & OpenAI")

# — Main form —
st.header("🖋️ Create Your Cover Letter")

# 1. Job description
job_text = st.text_area(
    "1. Job Description (URL or text)", 
    height=150
)

# 2. Resume uploader
uploaded_file = st.file_uploader(
    "2. Upload Your Resume (PDF or DOCX)", 
    type=["pdf", "docx"]
)

def extract_text(file) -> str:
    """Extracts and returns all text from a PDF or DOCX file."""
    fname = file.name.lower()
    if fname.endswith(".pdf"):
        reader = PdfReader(file)
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)
    elif fname.endswith(".docx"):
        doc = docx.Document(file)
        return "\n\n".join(p.text for p in doc.paragraphs)
    else:
        return ""

resume_text = ""
if uploaded_file:
    with st.spinner("Extracting text from resume…"):
        resume_text = extract_text(uploaded_file)
    st.success("✅ Resume loaded")
    # show first 300 chars as preview
    st.text_area("Resume Preview", resume_text, height=100)
else:
    st.info("Please upload a PDF or DOCX file to proceed.")

# 3. Your name & email
your_name = st.text_input("3. Your Full Name")
your_email = st.text_input("4. Your Email Address")

# 4. Style controls
col1, col2 = st.columns(2)
with col1:
    tone = st.selectbox(
        "Tone",
        ["Professional", "Friendly", "Enthusiastic", "Confident", "Concise"]
    )
    complexity = st.slider(
        "Complexity",
        min_value=1, max_value=10, value=5,
        help="1 = very simple wording, 10 = more elaborate phrasing"
    )
with col2:
    length = st.selectbox(
        "Letter Length",
        ["Short (1–2 paragraphs)", "Standard (3 paragraphs)", "Detailed (4–5 paragraphs)"]
    )
    include_greeting = st.checkbox(
        "Customize Greeting",
        value=False
    )
    if include_greeting:
        recipient_name = st.text_input("Recipient’s Name")
    else:
        recipient_name = None

# 5. Optional settings
with st.expander("Optional Settings"):
    company_name = st.text_input("Company Name")
    position_title = st.text_input("Position Title")
    closing_remark = st.text_input(
        "Custom Closing Phrase", 
        placeholder="e.g. Warm regards"
    )
    add_export = st.checkbox("Enable DOCX/PDF Export", value=False)

# 6. Generate button
if st.button("🚀 Generate Cover Letter"):
    if not (job_text and resume_text and your_name and your_email):
        st.error("Please fill out all required fields and upload your resume.")
    else:
        # Build prompt
        prompt_parts = [
            f"Write a {length.lower()} cover letter in a {tone.lower()} tone.",
            f"Complexity level: {complexity}/10.",
            f"Job description: {job_text}",
            f"My resume content:\n{resume_text}",
            f"My name is {your_name}, email: {your_email}."
        ]
        if company_name:
            prompt_parts.append(f"Address it to {company_name}.")
        if position_title:
            prompt_parts.append(f"Position title: {position_title}.")
        if recipient_name:
            prompt_parts.append(f"Use greeting: 'Dear {recipient_name},'")
        if closing_remark:
            prompt_parts.append(f"End with: '{closing_remark}, {your_name}.'")
        
        full_prompt = "\n".join(prompt_parts)

        # Call OpenAI
        with st.spinner("Generating your cover letter…"):
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[{"role":"user","content": full_prompt}],
                temperature=0.7,
                max_tokens=600,
            )
        letter = response.choices[0].message.content.strip()

        # Display result
        st.success("✅ Done!")
        st.text_area("Your Cover Letter", value=letter, height=300)

        # Export option
        if add_export:
            from io import BytesIO
            from docx import Document
            from docx.shared import Pt

            doc = Document()
            for line in letter.split("\n"):
                p = doc.add_paragraph(line)
                p.style.font.size = Pt(11)
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            st.download_button(
                "📄 Download as .docx",
                data=buffer,
                file_name="cover_letter.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
